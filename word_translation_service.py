import docx
import docx.shared
from docx.oxml.ns import qn
from docx.enum.text import WD_COLOR_INDEX
from docx.text.paragraph import Paragraph
import os
import copy
import threading
import concurrent.futures
import time
from typing import List, Dict, Tuple, Optional
import re
import logging
import asyncio

from translation import TranslationService

logger = logging.getLogger(__name__)

class WordTranslationService:
    """Word document translation service preserving format"""
    
    def __init__(self, api_key: str, base_url: str = "https://openrouter.ai/api/v1"):
        self.api_key = api_key
        self.base_url = base_url
        
        self.translator = TranslationService(api_key, base_url)
        
        # 并发配置
        self.MAX_WORKERS = 100
        self.RATE_LIMIT_DELAY = 0.1
        
    def copy_paragraph_format(self, source_paragraph, target_paragraph):
        """复制段落格式到目标段落"""
        try:
            # 复制段落级别的格式
            if source_paragraph.style:
                try:
                    target_paragraph.style = source_paragraph.style
                except:
                    pass
            
            # 复制对齐方式
            if source_paragraph.alignment is not None:
                target_paragraph.alignment = source_paragraph.alignment
            
            # 复制段落格式属性
            source_pf = source_paragraph.paragraph_format
            target_pf = target_paragraph.paragraph_format
            
            # 复制缩进
            if source_pf.left_indent is not None:
                target_pf.left_indent = source_pf.left_indent
            if source_pf.right_indent is not None:
                target_pf.right_indent = source_pf.right_indent
            if source_pf.first_line_indent is not None:
                target_pf.first_line_indent = source_pf.first_line_indent
            
            # 复制间距
            if source_pf.space_before is not None:
                target_pf.space_before = source_pf.space_before
            if source_pf.space_after is not None:
                target_pf.space_after = source_pf.space_after
            if source_pf.line_spacing is not None:
                target_pf.line_spacing = source_pf.line_spacing
            if source_pf.line_spacing_rule is not None:
                target_pf.line_spacing_rule = source_pf.line_spacing_rule
                
            # 复制其他格式
            if hasattr(source_pf, 'keep_together') and source_pf.keep_together is not None:
                target_pf.keep_together = source_pf.keep_together
            if hasattr(source_pf, 'keep_with_next') and source_pf.keep_with_next is not None:
                target_pf.keep_with_next = source_pf.keep_with_next
            if hasattr(source_pf, 'page_break_before') and source_pf.page_break_before is not None:
                target_pf.page_break_before = source_pf.page_break_before
            if hasattr(source_pf, 'widow_control') and source_pf.widow_control is not None:
                target_pf.widow_control = source_pf.widow_control
                
        except Exception as e:
            logger.error(f"复制段落格式时出错: {e}")

    def copy_run_format(self, source_run, target_run, override_color=True):
        """复制运行格式到目标运行"""
        try:
            source_font = source_run.font
            target_font = target_run.font
            
            # 复制字体属性
            if source_font.name is not None:
                target_font.name = source_font.name
            if source_font.size is not None:
                target_font.size = source_font.size
            if source_font.bold is not None:
                target_font.bold = source_font.bold
            if source_font.italic is not None:
                target_font.italic = source_font.italic
            if source_font.underline is not None:
                target_font.underline = source_font.underline
            if source_font.strike is not None:
                target_font.strike = source_font.strike
            if source_font.subscript is not None:
                target_font.subscript = source_font.subscript
            if source_font.superscript is not None:
                target_font.superscript = source_font.superscript
            if source_font.all_caps is not None:
                target_font.all_caps = source_font.all_caps
            if source_font.small_caps is not None:
                target_font.small_caps = source_font.small_caps
            if source_font.shadow is not None:
                target_font.shadow = source_font.shadow
            if source_font.emboss is not None:
                target_font.emboss = source_font.emboss
            if source_font.imprint is not None:
                target_font.imprint = source_font.imprint
            if source_font.outline is not None:
                target_font.outline = source_font.outline
            
            # 设置颜色
            if override_color:
                target_font.color.rgb = docx.shared.RGBColor(255, 0, 0)  # 红色
            elif source_font.color.rgb is not None:
                target_font.color.rgb = source_font.color.rgb
                
            # 复制高亮
            if hasattr(source_font, 'highlight_color') and source_font.highlight_color is not None:
                target_font.highlight_color = source_font.highlight_color
                
        except Exception as e:
            logger.error(f"复制运行格式时出错: {e}")

    def replace_paragraph_text_keep_format(self, paragraph, new_text: str) -> bool:
        """替换段落文本但保持格式和图像"""
        try:
            # 保存图像runs（带位置信息）
            image_runs_info = []
            for i, run in enumerate(paragraph.runs):
                has_image = bool(run._element.findall('.//w:drawing', namespaces=run._element.nsmap)) or \
                           bool(run._element.findall('.//w:pict', namespaces=run._element.nsmap))
                if has_image:
                    image_runs_info.append((i, copy.deepcopy(run._element)))
            
            # 保存第一个文本run的格式
            first_text_run = None
            for run in paragraph.runs:
                has_image = bool(run._element.findall('.//w:drawing', namespaces=run._element.nsmap)) or \
                           bool(run._element.findall('.//w:pict', namespaces=run._element.nsmap))
                if not has_image:
                    first_text_run = run
                    break
            
            # 清空所有run的文本内容（但保留图像）
            for run in paragraph.runs:
                has_image = bool(run._element.findall('.//w:drawing', namespaces=run._element.nsmap)) or \
                           bool(run._element.findall('.//w:pict', namespaces=run._element.nsmap))
                if not has_image:
                    run.text = ''
            
            # 在第一个文本run中设置新文本
            if first_text_run is not None:
                first_text_run.text = new_text
                # 确保颜色是黑色
                first_text_run.font.color.rgb = docx.shared.RGBColor(0, 0, 0)
            elif paragraph.runs:
                # 如果没有文本run，在第一个run中设置
                paragraph.runs[0].text = new_text
                paragraph.runs[0].font.color.rgb = docx.shared.RGBColor(0, 0, 0)
            else:
                # 如果没有任何run，创建一个新的
                run = paragraph.add_run(new_text)
                run.font.color.rgb = docx.shared.RGBColor(0, 0, 0)
            
            return True
        except Exception as e:
            logger.error(f"替换段落文本时出错: {e}")
            return False

    def insert_translation_simple(self, paragraph, translated_text: str) -> bool:
        """简单的翻译插入方法。返回新插入的段落对象，失败返回 False。"""
        try:
            # 获取段落所在的父元素
            parent = paragraph._element.getparent()
            
            # 创建新的段落元素
            new_para = docx.oxml.parse_xml(r'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
            
            # 创建段落属性元素来复制格式
            if paragraph._element.find(qn('w:pPr')) is not None:
                original_ppr = paragraph._element.find(qn('w:pPr'))
                new_ppr = copy.deepcopy(original_ppr)
                new_para.insert(0, new_ppr)
            
            # 创建运行元素并设置文本
            run = docx.oxml.parse_xml(r'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
            text_elem = docx.oxml.parse_xml(r'<w:t xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
            text_elem.text = translated_text
            run.append(text_elem)
            
            # 设置运行格式
            rpr = docx.oxml.parse_xml(r'<w:rPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
            
            # 复制原始运行的格式属性（如果存在）
            if paragraph.runs:
                original_run = paragraph.runs[0]._element
                original_rpr = original_run.find(qn('w:rPr'))
                if original_rpr is not None:
                    # 复制原始格式
                    new_rpr = copy.deepcopy(original_rpr)
                    # 更新颜色为红色
                    color_elem = new_rpr.find(qn('w:color'))
                    if color_elem is not None:
                        color_elem.set(qn('w:val'), 'FF0000')
                    else:
                        color = docx.oxml.parse_xml(r'<w:color xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="FF0000"/>')
                        new_rpr.append(color)
                    rpr = new_rpr
                else:
                    # 只设置红色
                    color = docx.oxml.parse_xml(r'<w:color xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="FF0000"/>')
                    rpr.append(color)
            else:
                # 只设置红色
                color = docx.oxml.parse_xml(r'<w:color xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="FF0000"/>')
                rpr.append(color)
            
            run.insert(0, rpr)
            new_para.append(run)
            
            # 在原段落后插入新段落
            parent.insert(list(parent).index(paragraph._element) + 1, new_para)
            # 返回段落对象
            try:
                new_paragraph = Paragraph(new_para, paragraph._parent)
            except Exception:
                # 回退：无法构建段落对象时返回 True 但不提供对象
                return True
            return new_paragraph
            
        except Exception as e:
            logger.error(f"插入翻译失败: {e}")
            return False

    async def process_document_dual_output(self, file_path: str, contrast_output_path: str, 
                                   translation_only_output_path: str,
                                   source_language: str = "English",
                                   target_language: str = "Chinese") -> List[Dict]:
        """处理文档并生成两个输出：对照翻译和仅译文"""
 
        # 读取原始文档
        original_doc = docx.Document(file_path)
        
        # 收集所有需要翻译的内容
        to_translate = []  # (type, element_info, text)
        
        # 收集段落
        for i, paragraph in enumerate(original_doc.paragraphs):
            text = paragraph.text.strip()
            if text:
                to_translate.append(('paragraph', i, text))
        
        # 收集表格内容
        table_cells = []
        for table_idx, table in enumerate(original_doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, para in enumerate(cell.paragraphs):
                        cell_text = para.text.strip()
                        if cell_text:
                            to_translate.append(('table_cell', (table_idx, row_idx, cell_idx, para_idx), cell_text))
        
        if not to_translate:
            return []
                    
        # 翻译所有文本
        texts = [item[2] for item in to_translate]
        translated_results = await self.translator.translate_texts_parallel(texts, source_language, target_language)
        
        # 生成仅译文文档
        translation_only_doc = docx.Document(file_path)
        
        # 替换仅译文文档的内容
        for item, tr in zip(to_translate, translated_results):
            translated_text, _references = tr
            typ, info, orig = item
            if typ == 'paragraph':
                paragraph_idx = info
                if paragraph_idx < len(translation_only_doc.paragraphs):
                    self.replace_paragraph_text_keep_format(translation_only_doc.paragraphs[paragraph_idx], translated_text)
            elif typ == 'table_cell':
                table_idx, row_idx, cell_idx, para_idx = info
                if (table_idx < len(translation_only_doc.tables) and 
                    row_idx < len(translation_only_doc.tables[table_idx].rows) and
                    cell_idx < len(translation_only_doc.tables[table_idx].rows[row_idx].cells) and
                    para_idx < len(translation_only_doc.tables[table_idx].rows[row_idx].cells[cell_idx].paragraphs)):
                    para = translation_only_doc.tables[table_idx].rows[row_idx].cells[cell_idx].paragraphs[para_idx]
                    self.replace_paragraph_text_keep_format(para, translated_text)
        
        # 保存仅译文文档
        translation_only_doc.save(translation_only_output_path)
        
        # 生成对照翻译文档
        contrast_doc = docx.Document(file_path)
        translated_paragraphs = []
        
        # 按倒序插入译文（避免索引变化）
        for item, tr in zip(reversed(to_translate), reversed(translated_results)):
            translated_text, references = tr
            typ, info, orig = item
            if typ == 'paragraph':
                paragraph_idx = info
                if paragraph_idx < len(contrast_doc.paragraphs):
                    original_para = contrast_doc.paragraphs[paragraph_idx]
                    inserted_para = self.insert_translation_simple(original_para, translated_text)
                    if inserted_para:
                        # 获取术语的source_type信息
                        from extraction_term import find_text_in_db
                        from translation import TranslationService
                        translation_service = TranslationService("", "")
                        _, source_types = find_text_in_db(orig, src_lang=source_language)
                        
                        # 需要高亮的原文术语（仅用户上传的）
                        user_terms = [term for term in references.keys() if source_types.get(term) == 'usr']
                        if user_terms:
                            self.highlight_terms_by_run(original_para, user_terms)

                        if isinstance(inserted_para, Paragraph):
                            # 需要高亮的译文术语（仅用户上传的）
                            user_translated_terms = [references[term] for term in user_terms]
                            if user_translated_terms:
                                self.highlight_terms_by_run(inserted_para, user_translated_terms)
                        translated_paragraphs.append({'original': orig, 'translated': translated_text})
            elif typ == 'table_cell':
                table_idx, row_idx, cell_idx, para_idx = info
                if (table_idx < len(contrast_doc.tables) and 
                    row_idx < len(contrast_doc.tables[table_idx].rows) and
                    cell_idx < len(contrast_doc.tables[table_idx].rows[row_idx].cells)):
                    cell = contrast_doc.tables[table_idx].rows[row_idx].cells[cell_idx]
                    # 在表格单元格中添加译文段落
                    trans_para = cell.add_paragraph(translated_text)
                    for run in trans_para.runs:
                        run.font.color.rgb = docx.shared.RGBColor(255, 0, 0)
                    
                    # 获取术语的source_type信息
                    from extraction_term import find_text_in_db
                    from translation import TranslationService
                    translation_service = TranslationService("", "")
                    _, source_types = find_text_in_db(orig, src_lang=source_language)
                    
                    # 高亮原文单元格中的术语（仅用户上传的）
                    if para_idx < len(cell.paragraphs):
                        para_obj = cell.paragraphs[para_idx]
                        user_terms = [term for term in references.keys() if source_types.get(term) == 'usr']
                        if user_terms:
                            self.highlight_terms_by_run(para_obj, user_terms)
            
                    translated_paragraphs.append({'original': orig, 'translated': translated_text})
        
        translated_paragraphs.reverse()
        
        # 保存对照翻译文档
        contrast_doc.save(contrast_output_path)
        
        return translated_paragraphs


    def highlight_terms_by_run(self, paragraph, terms: list[str], case_insensitive: bool = True) -> None:
        """Precisely highlight glossary terms inside a paragraph with yellow color."""
        if not terms:
            return
        normalized_terms = [t.lower() for t in terms if isinstance(t, str) and t.strip()]
        if not normalized_terms:
            return
        # Sort by length to prefer longer terms when alternatives overlap
        normalized_terms.sort(key=len, reverse=True)
        pattern_text = "|".join(re.escape(t) for t in normalized_terms)
        flags = re.IGNORECASE if case_insensitive else 0
        try:
            term_pattern = re.compile(f"({pattern_text})", flags)
        except re.error:
            # Fallback: if regex compilation fails for any reason, do nothing
            return
        # Work on a copy because we're going to mutate paragraph runs
        for run in list(paragraph.runs):
            text = (run.text or "").lower()
            matches = list(term_pattern.finditer(text))
            if not matches:
                run.font.highlight_color = None
            segments: List[Tuple[str, bool]] = []  # (text, should_highlight)
            cursor = 0
            for m in matches:
                if m.start() > cursor:
                    segments.append((text[cursor:m.start()], False))
                segments.append((text[m.start():m.end()], True))
                cursor = m.end()
            if cursor < len(text):
                segments.append((text[cursor:], False))

            # Replace the original run with the first segment
            first_text, should_highlight = segments[0]
            run.text = first_text
            # Clear any existing highlight first
            try:
                run.font.highlight_color = None
            except Exception:
                pass
            if should_highlight:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW

            prev_r = run._element

            # Insert subsequent segments as new runs placed after the current one
            for seg_text, should_highlight in segments[1:]:
                # Clone the original run's XML so we keep formatting
                new_r = copy.deepcopy(prev_r)

                # Ensure there's a text node and set text
                t = new_r.find(qn('w:t'))
                if t is None:
                    t = docx.oxml.parse_xml(r'<w:t xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
                    new_r.append(t)
                t.text = seg_text

                # Ensure rPr exists and set/remove highlight
                rPr = new_r.find(qn('w:rPr'))
                if rPr is None:
                    rPr = docx.oxml.parse_xml(r'<w:rPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
                    new_r.insert(0, rPr)
                existing_hl = rPr.find(qn('w:highlight'))
                if existing_hl is not None:
                    rPr.remove(existing_hl)
                if should_highlight:
                    hl = docx.oxml.parse_xml(r'<w:highlight xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="yellow"/>')
                    rPr.append(hl)

                prev_r.addnext(new_r)
                prev_r = new_r
            
    
    async def extract_and_translate_doc(self, file_path: str, contrast_output_path: str, 
                                translation_only_output_path: str,
                                source_language: str = "English",
                                target_language: str = "Chinese") -> List[Dict]:
        """从doc文件中提取文本并生成两个翻译文档"""
        try:
            # 对于.doc文件，先提取文本然后创建带翻译的docx
            import docx2txt
            text = docx2txt.process(file_path)
            paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
            
            if not paragraphs:
                return []
            
            # 并行翻译所有段落
            translated_texts = await self.translator.translate_texts_parallel(paragraphs, source_language, target_language)
            
            # 创建对照翻译文档
            contrast_doc = docx.Document()
            # 创建仅译文文档
            translation_only_doc = docx.Document()
            
            translated_paragraphs = []
            
            for paragraph_text, translated_text in zip(paragraphs, translated_texts):
                # 对照文档：原文 + 译文
                original_para = contrast_doc.add_paragraph(paragraph_text)
                translated_para = contrast_doc.add_paragraph()
                self.copy_paragraph_format(original_para, translated_para)
                translated_run = translated_para.add_run(translated_text)
                translated_run.font.color.rgb = docx.shared.RGBColor(255, 0, 0)  # 红色
                
                # 仅译文文档：只有译文
                trans_only_para = translation_only_doc.add_paragraph(translated_text)
                
                translated_paragraphs.append({
                    'original': paragraph_text,
                    'translated': translated_text
                })
            
            # 保存两个文档
            contrast_doc.save(contrast_output_path)
            translation_only_doc.save(translation_only_output_path)
            
            return translated_paragraphs
        except ImportError:
            raise Exception("处理.doc文件需要安装docx2txt库: pip install docx2txt")
        except Exception as e:
            raise Exception(f"处理doc文件失败: {str(e)}") 